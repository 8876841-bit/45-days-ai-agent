from __future__ import annotations

from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "learning-log" / "day01-45-detailed-learning-summaries.md"
OUT = ROOT / "learning-log" / "day01-45-detailed-learning-summaries.docx"


def set_east_asia_font(run, font_name: str) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)

    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Title", 24, "0B2545", 0, 12),
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    code = styles.add_style("Code Block", 1)
    code.font.name = "Consolas"
    code._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    code.font.size = Pt(9.5)
    code.font.color.rgb = RGBColor(50, 50, 50)
    code.paragraph_format.left_indent = Inches(0.22)
    code.paragraph_format.right_indent = Inches(0.1)
    code.paragraph_format.space_before = Pt(4)
    code.paragraph_format.space_after = Pt(6)
    code.paragraph_format.line_spacing = 1.15


def add_cover(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.style = doc.styles["Title"]
    run = title.add_run("Day 1-45 详细学习心得复述版")
    run.bold = True
    set_east_asia_font(run, "Microsoft YaHei")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("45 天 AI Agent 学习计划复习文档")
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(85, 85, 85)
    set_east_asia_font(run, "Microsoft YaHei")

    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=3)
    table.autofit = False
    widths = [Inches(1.8), Inches(2.3), Inches(2.4)]
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    labels = ["用途", "覆盖范围", "阅读方式"]
    for cell, label, width in zip(hdr.cells, labels, widths):
        cell.width = width
        set_cell_shading(cell, "E8EEF5")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(label)
        r.bold = True
        set_east_asia_font(r, "Microsoft YaHei")

    row = table.add_row()
    values = [
        "复习与跨设备延续学习",
        "Day 1 到 Day 45",
        "每天按主题阅读，重点看“关键理解”和“保留结论”",
    ]
    for cell, value, width in zip(row.cells, values, widths):
        cell.width = width
        p = cell.paragraphs[0]
        r = p.add_run(value)
        set_east_asia_font(r, "Microsoft YaHei")

    doc.add_paragraph()
    note = doc.add_paragraph()
    r = note.add_run(
        "说明：这份文档不是短视频口播稿，而是用于复习的详细复述版。"
        "重点是把每天学到的概念用自己的话讲清楚。"
    )
    r.italic = True
    r.font.color.rgb = RGBColor(85, 85, 85)
    set_east_asia_font(r, "Microsoft YaHei")

    doc.add_section(WD_SECTION.NEW_PAGE)


def add_footer(doc: Document) -> None:
    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("45 天 AI Agent 学习计划 · 详细学习心得")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(100, 100, 100)
        set_east_asia_font(run, "Microsoft YaHei")


def add_code_paragraph(doc: Document, line: str) -> None:
    p = doc.add_paragraph(style="Code Block")
    r = p.add_run(line if line else " ")
    set_east_asia_font(r, "Microsoft YaHei")


def add_body_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_east_asia_font(r, "Microsoft YaHei")


def build() -> None:
    text = SOURCE.read_text(encoding="utf-8")
    doc = Document()
    configure_document(doc)
    add_cover(doc)

    in_code = False
    code_lines: list[str] = []

    for raw_line in text.splitlines():
        line = raw_line.rstrip()

        if line.startswith("```"):
            if in_code:
                for code_line in code_lines:
                    add_code_paragraph(doc, code_line)
                code_lines = []
                in_code = False
            else:
                in_code = True
            continue

        if in_code:
            code_lines.append(line)
            continue

        if not line.strip():
            continue

        if line == "---":
            continue

        if line.startswith("# "):
            p = doc.add_paragraph(style="Title")
            r = p.add_run(line[2:].strip())
            r.bold = True
            set_east_asia_font(r, "Microsoft YaHei")
            continue

        if line.startswith("## "):
            title = line[3:].strip()
            if re.match(r"Day\s+\d+", title):
                # Start each day on a new page except the first content section after cover.
                if len(doc.paragraphs) > 3:
                    doc.add_page_break()
            p = doc.add_paragraph(style="Heading 1")
            r = p.add_run(title)
            r.bold = True
            set_east_asia_font(r, "Microsoft YaHei")
            continue

        if line.startswith("### "):
            p = doc.add_paragraph(style="Heading 2")
            r = p.add_run(line[4:].strip())
            r.bold = True
            set_east_asia_font(r, "Microsoft YaHei")
            continue

        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run(line[2:].strip())
            set_east_asia_font(r, "Microsoft YaHei")
            continue

        add_body_paragraph(doc, line)

    add_footer(doc)
    doc.save(OUT)


if __name__ == "__main__":
    build()
